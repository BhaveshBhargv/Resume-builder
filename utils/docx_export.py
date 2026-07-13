"""Export a ResumeData object to an ATS-friendly Word (.docx) document.

The layout is deliberately simple: a single column, standard fonts, real
Word bullet lists, and plain uppercase section headings with a thin
underline. No tables, text boxes, columns, images, or graphics -- those are
exactly the things applicant-tracking systems fail to parse. This mirrors the
whole app's ATS-first philosophy.
"""
from __future__ import annotations

from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt

from models.resume_data import ResumeData


def _date_range(start: str, end: str, is_current: bool) -> str:
    end_text = "Present" if is_current else end
    if start and end_text:
        return f"{start} – {end_text}"
    return start or end_text or ""


def _add_bottom_border(paragraph) -> None:
    """Draw a thin rule under a heading paragraph (kept text-only for ATS)."""
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "888888")
    borders.append(bottom)
    p_pr.append(borders)


def build_docx(resume: ResumeData) -> bytes:
    """Render the resume to .docx bytes suitable for st.download_button."""
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Pt(40)
        section.left_margin = section.right_margin = Pt(50)

    pi = resume.personal_info

    def heading(text: str) -> None:
        para = doc.add_paragraph()
        para.space_before = Pt(8)
        run = para.add_run(text.upper())
        run.bold = True
        run.font.size = Pt(12)
        _add_bottom_border(para)

    # --- Header: name + contact line -----------------------------------------
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run(pi.full_name or "Your Name")
    name_run.bold = True
    name_run.font.size = Pt(20)

    contact = " | ".join(
        bit for bit in [pi.email, pi.phone, pi.location, pi.linkedin_url, pi.portfolio_url] if bit
    )
    if contact:
        c = doc.add_paragraph()
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        c.add_run(contact).font.size = Pt(9)

    # --- Professional summary -------------------------------------------------
    if pi.professional_summary:
        heading("Professional Summary")
        doc.add_paragraph(pi.professional_summary)

    # --- Experience -----------------------------------------------------------
    if resume.experience:
        heading("Experience")
        for exp in resume.experience:
            line = doc.add_paragraph()
            title = ", ".join(x for x in [exp.job_title, exp.company] if x)
            line.add_run(title).bold = True
            meta = " | ".join(x for x in [exp.location, _date_range(exp.start_date, exp.end_date, exp.is_current)] if x)
            if meta:
                line.add_run(f"  —  {meta}")
            for bullet in exp.bullet_points:
                doc.add_paragraph(bullet, style="List Bullet")

    # --- Education ------------------------------------------------------------
    if resume.education:
        heading("Education")
        for edu in resume.education:
            line = doc.add_paragraph()
            degree = ", ".join(x for x in [edu.degree, edu.field_of_study] if x)
            title = " — ".join(x for x in [degree, edu.institution] if x)
            line.add_run(title).bold = True
            extras = [_date_range(edu.start_date, edu.end_date, edu.is_current)]
            if edu.gpa:
                extras.append(f"GPA: {edu.gpa}")
            meta = " | ".join(x for x in extras if x)
            if meta:
                line.add_run(f"  —  {meta}")
            for ach in edu.achievements:
                doc.add_paragraph(ach, style="List Bullet")

    # --- Projects -------------------------------------------------------------
    if resume.projects:
        heading("Projects")
        for proj in resume.projects:
            line = doc.add_paragraph()
            line.add_run(proj.name or "Project").bold = True
            if proj.technologies:
                line.add_run(f"  —  {', '.join(proj.technologies)}")
            if proj.description:
                doc.add_paragraph(proj.description)
            for bullet in proj.bullet_points:
                doc.add_paragraph(bullet, style="List Bullet")
            if proj.url:
                doc.add_paragraph(proj.url)

    # --- Skills ---------------------------------------------------------------
    if any(cat.skills for cat in resume.skills):
        heading("Skills")
        for cat in resume.skills:
            if cat.skills:
                para = doc.add_paragraph()
                para.add_run(f"{cat.category_name}: ").bold = True
                para.add_run(", ".join(cat.skills))

    # --- Extra sections (from an uploaded resume) -----------------------------
    for extra in resume.extra_sections:
        if extra.heading or extra.content:
            heading(extra.heading or "Additional Information")
            if extra.content:
                doc.add_paragraph(extra.content)

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
